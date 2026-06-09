from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import openpyxl


ROOT = Path(__file__).resolve().parents[2]
PROJECT_ROOT = ROOT.parent

TEMPLATE_DOCX = PROJECT_ROOT / "CO2_Tissue O2_Supplemental Digital Content_R1_4_19_2026.docx"
ETABLE_1_2 = ROOT / "docs" / "manuscript_development" / "generated_assets" / "supplementary_etable1_2_cohort_characteristics.xlsx"
ETABLE_3_5 = ROOT / "results" / "supplemental_etables" / "Supplemental_eTables3_5_CO2_rSO2.xlsx"
ETABLE_6_8 = ROOT / "results" / "supplemental_etables" / "Supplemental_eTables6_8_CO2_rSO2.xlsx"
OUT_DOCX = ROOT / "docs" / "manuscript_development" / "generated_assets" / "CO2_Tissue_O2_Supplemental_Digital_Content_integrated_eTables1_8_20260609.docx"


TITLE = "End-tidal carbon dioxide is associated with cerebral tissue oxygenation during off-pump coronary bypass"

COHORT_LABELS = {
    "rSO2_Ch1": "Left SctO2 cohort",
    "rSO2_Ch2": "Right SctO2 cohort",
    "rSO2_Ch3": "SftO2 cohort",
}

STAGE_LABELS = {
    "raw_timeseries_rows": "Raw time-series observations",
    "after_required_etco2_y_nonmissing": "After requiring nonmissing EtCO2 and site-specific tissue oxygenation",
    "after_cohort_clip_to_missing_and_dropna": "After cohort clipping to required nonmissing variables",
    "final_usable_points_strict_etco2_rso2": "Final usable observations after strict EtCO2 and tissue oxygenation ranges",
    "after_intraop_covars_clip_to_missing": "After intraoperative covariate range clipping",
    "after_intraop_covars_fill": "After intraoperative covariate imputation",
}

TABLE_LABELS = {
    "Table 1 baseline/static": "Baseline and patient-level characteristics",
    "Table 2 intraoperative/timestamp": "Intraoperative timestamp-level characteristics",
}

CHAR_LABELS = {
    "Cardiac_index": "Cardiac index",
    "Carotid_artery_disease": "Carotid artery disease",
    "Diabetes_status": "Diabetes",
    "Drinking_status": "Drinking history",
    "ET_CO2": "EtCO2",
    "FiO2_new": "FiO2",
    "Hb": "Hemoglobin",
    "Hypertension_140_90": "Preoperative hypertension >140/90 mmHg",
    "Left_SctO2": "Left SctO2",
    "Mean_blood_pressure": "Mean blood pressure",
    "Right_SctO2": "Right SctO2",
    "Smoking_new": "Smoking history",
    "SstO2": "SftO2",
    "Statin_1": "Statin use",
    "TEMP": "Temperature",
    "rSO2_Ch1": "Left SctO2",
    "rSO2_Ch2": "Right SctO2",
    "rSO2_Ch3": "SftO2",
}


def text_value(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return f"{int(value):,}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def prettify_label(text):
    if text is None:
        return ""
    out = str(text)
    for old, new in sorted(CHAR_LABELS.items(), key=lambda x: len(x[0]), reverse=True):
        out = out.replace(old, new)
    out = out.replace("=0.0", " = 0")
    out = out.replace("=1.0", " = 1")
    out = out.replace("=2.0", " = 2")
    out = out.replace("=3.0", " = 3")
    return out


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_cell_text(cell, value, bold=False, size=7.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text_value(value))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")


def add_paragraph(doc, text, style=None, bold=False, size=9):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(13)
    run.bold = True
    return paragraph


def add_table(doc, rows, header_rows=1, font_size=7.5, group_rows=None):
    group_rows = set(group_rows or [])
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Normal Table"
    except KeyError:
        pass
    set_table_borders(table)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            is_header = r_idx < header_rows
            is_group = r_idx in group_rows
            set_cell_text(table.cell(r_idx, c_idx), value, bold=is_header or is_group, size=font_size)
            if is_header:
                shade_cell(table.cell(r_idx, c_idx), "EDEDED")
            if is_group:
                shade_cell(table.cell(r_idx, c_idx), "F6F6F6")
        if r_idx < header_rows:
            set_repeat_table_header(table.rows[r_idx])
    add_paragraph(doc, "", size=1)
    return table


def load_workbook(path):
    return openpyxl.load_workbook(path, data_only=True, read_only=True)


def build_etable1_rows():
    wb = load_workbook(ETABLE_1_2)
    ws = wb["flow_counts"]
    records = list(ws.iter_rows(min_row=2, values_only=True))
    stages = []
    for _, stage, _, _ in records:
        if stage not in stages:
            stages.append(stage)
    rows = [["Filtering stage", "Left SctO2 rows, n", "Left SctO2 patients, n", "Right SctO2 rows, n", "Right SctO2 patients, n", "SftO2 rows, n", "SftO2 patients, n"]]
    by_key = {(ycol, stage): (n_rows, n_patients) for ycol, stage, n_rows, n_patients in records}
    for stage in stages:
        row = [STAGE_LABELS.get(stage, stage)]
        for ycol in ("rSO2_Ch1", "rSO2_Ch2", "rSO2_Ch3"):
            n_rows, n_patients = by_key[(ycol, stage)]
            row.extend([n_rows, n_patients])
        rows.append(row)
    return rows


def build_etable2_rows():
    wb = load_workbook(ETABLE_1_2)
    ws = wb["wide"]
    source = list(ws.iter_rows(min_row=2, values_only=True))
    rows = [["Characteristic", "Left SctO2", "Right SctO2", "SftO2", "Missing Left SctO2, n", "Missing Right SctO2, n", "Missing SftO2, n"]]
    group_rows = []
    current_group = None
    for record in source:
        group = record[0]
        if group != current_group:
            current_group = group
            group_rows.append(len(rows))
            rows.append([TABLE_LABELS.get(group, group), "", "", "", "", "", ""])
        label = prettify_label(record[1])
        if group == "Table 2 intraoperative/timestamp" and label == "Patients, n":
            label = "Timestamp observations, n"
        rows.append([label] + [text_value(v) for v in record[2:]])
    return rows, group_rows


def split_sheet(ws):
    raw_rows = list(ws.iter_rows(values_only=True))
    title = text_value(raw_rows[0][0])
    table_rows = []
    notes = []
    seen_header = False
    in_notes = False
    for row in raw_rows[2:]:
        non_empty = any(cell is not None for cell in row)
        if not non_empty:
            if seen_header:
                in_notes = True
            continue
        if in_notes:
            notes.append(text_value(row[0]))
        else:
            seen_header = True
            table_rows.append([prettify_label(text_value(cell)) for cell in row])
    return title, table_rows, notes


def add_notes(doc, notes):
    for note in notes:
        if note:
            add_paragraph(doc, note, size=8)


def main():
    doc = Document(TEMPLATE_DOCX)
    clear_body(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    add_paragraph(doc, "Supplemental Digital Content", bold=True, size=11)
    add_paragraph(doc, "The authors provided the following additional information for the manuscript.", size=9)
    add_paragraph(doc, TITLE, size=9)

    add_heading(doc, "eTable 1. Cohort flow and signal availability in the left cerebral, right cerebral, and forearm tissue oxygenation analytic cohorts")
    add_table(doc, build_etable1_rows(), font_size=6.5)
    add_notes(doc, [
        "SctO2, cerebral tissue oxygen saturation; SftO2, forearm tissue oxygen saturation; EtCO2, end-tidal carbon dioxide.",
        "Rows represent synchronized timestamp-level observations unless otherwise indicated.",
    ])

    add_heading(doc, "eTable 2. Baseline and intraoperative characteristics in the left cerebral, right cerebral, and forearm tissue oxygenation analytic cohorts")
    rows, group_rows = build_etable2_rows()
    add_table(doc, rows, font_size=6.2, group_rows=group_rows)
    add_notes(doc, [
        "SctO2, cerebral tissue oxygen saturation; SftO2, forearm tissue oxygen saturation; EtCO2, end-tidal carbon dioxide; FiO2, fraction of inspired oxygen.",
        "Continuous variables are presented as mean (SD) or median (IQR), and categorical variables as n (%). Missing values are shown as n.",
    ])

    for path, sheet_names in (
        (ETABLE_3_5, ["eTable3_artifact", "eTable4_imputation", "eTable5_patient_level"]),
        (ETABLE_6_8, ["eTable6_model_fit", "eTable7_continuous", "eTable8_categorical"]),
    ):
        wb = load_workbook(path)
        for sheet_name in sheet_names:
            title, rows, notes = split_sheet(wb[sheet_name])
            add_heading(doc, title)
            font_size = 6.3 if len(rows[0]) >= 7 else 7.0
            add_table(doc, rows, font_size=font_size)
            add_notes(doc, notes)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
