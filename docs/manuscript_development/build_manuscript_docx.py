from pathlib import Path
import re
import argparse
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
REPO_ROOT = ROOT.parents[1]
DEFAULT_SOURCE = ROOT / "CO2_MANUSCRIPT_DRAFT_V1.md"
DEFAULT_OUTPUT = ROOT / "CO2_MANUSCRIPT_DRAFT_V1.docx"
ASSET_DIR = REPO_ROOT / "outputs_local" / "docs" / "manuscript_development" / "generated_assets"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = getattr(tbl_pr, "tblW", None)
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)


def set_cell_borders(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in edges.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_runs(paragraph, text):
    # Minimal inline Markdown support for `code` and **bold**.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in row):
            rows.append(row)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
            cell.text = value
    doc.add_paragraph()


def read_csv_rows(path):
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.reader(f))


def add_caption(doc, label, title, legend=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label} | {title}")
    run.bold = True
    if legend:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_runs(p, legend)


def add_nature_table(doc, label, title, legend, csv_name):
    rows = read_csv_rows(ASSET_DIR / csv_name)
    add_caption(doc, label, title, legend)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_width(table)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_margins(cell, top=60, start=90, bottom=60, end=90)
            cell.text = value
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(7)
                    if r_idx == 0:
                        run.bold = True
            if r_idx == 0:
                set_cell_borders(
                    cell,
                    top={"val": "single", "sz": "8", "color": "000000"},
                    bottom={"val": "single", "sz": "8", "color": "000000"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
            elif r_idx == len(rows) - 1:
                set_cell_borders(
                    cell,
                    bottom={"val": "single", "sz": "8", "color": "000000"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
            else:
                set_cell_borders(
                    cell,
                    top={"val": "nil"},
                    bottom={"val": "nil"},
                    left={"val": "nil"},
                    right={"val": "nil"},
                )
    doc.add_paragraph()


def add_figure(doc, label, title, image_name, legend):
    add_caption(doc, label, title, legend)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSET_DIR / image_name), width=Inches(6.2))
    doc.add_paragraph()


def handle_display_placeholder(doc, line):
    if line.startswith("[Insert Table 1"):
        add_nature_table(
            doc,
            "Table 1",
            "Cohort characteristics.",
            "Values are shown by channel-specific analytic cohort. Continuous variables are mean (s.d.); "
            "categorical variables are n (%). SctO2, cerebral tissue oxygen saturation; "
            "SftO2, forearm tissue oxygen saturation.",
            "table1_cohort_characteristics.csv",
        )
        return True
    if line.startswith("[Insert Figure 1"):
        add_figure(
            doc,
            "Figure 1",
            "Cohort assembly by tissue oxygenation channel.",
            "figure1_cohort_flow.png",
            "Flow counts show patient and timestamp retention for left SctO2, right SctO2, and SftO2 analytic cohorts.",
        )
        return True
    if line.startswith("[Insert Table 2"):
        add_nature_table(
            doc,
            "Table 2",
            "Adjusted clinical-step contrasts.",
            "Contrasts are model-based adjusted differences in rSO2 for prespecified exposure increments. "
            "Confidence intervals are derived from archived prediction matrices and should be interpreted "
            "as model-based uncertainty intervals.",
            "table2_clinical_step_contrasts.csv",
        )
        add_figure(
            doc,
            "Figure 3",
            "Clinical-step comparison of EtCO2, FiO2, and temperature.",
            "figure3_clinical_step_contrasts.png",
            "Points show adjusted differences in rSO2 for clinically interpretable exposure increments; horizontal bars show 95% intervals.",
        )
        return True
    if line.startswith("[Insert Figure 2"):
        add_figure(
            doc,
            "Figure 2",
            "Adjusted EtCO2-rSO2 response curves.",
            "figure2_etco2_adjusted_curves.png",
            "Curves show adjusted tissue oxygenation across the observed EtCO2 range by channel. Bands show uncertainty intervals derived from archived prediction matrices. Bottom ticks show modeled curve support, not the observed EtCO2 histogram.",
        )
        return True
    if line.startswith("[Insert Supplementary Figure 1"):
        add_figure(
            doc,
            "Supplementary Figure 1",
            "Local slopes of the adjusted EtCO2-rSO2 relationship.",
            "figure4_etco2_local_slopes.png",
            "Local slopes summarize descriptive model gradients across EtCO2 mmHg bins from 20 to 50 mmHg. These summaries are descriptive and do not define treatment thresholds.",
        )
        return True
    return False


def build(source=DEFAULT_SOURCE, output=DEFAULT_OUTPUT):
    doc = Document()
    style_document(doc)
    lines = source.read_text(encoding="utf-8").splitlines()

    title_done = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if handle_display_placeholder(doc, line):
            i += 1
            continue
        if line.startswith("# "):
            text = line[2:].strip()
            if not title_done:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor.from_string("0B2545")
                title_done = True
            else:
                doc.add_heading(text, level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
        i += 1

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CO2-rSO2 manuscript draft V1")

    doc.save(output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.source, args.output)
